import os
from random import randint
from docx import Document
from docx.shared import Pt
import printa,gui
lis=gui.run()
doc=Document()
p=doc.add_paragraph("")
abs_path = os.path.abspath(__file__)
abs_path=abs_path.replace("index.py","")
index=0
while index <=lis[0]-1:
    a=randint(1,lis[1])
    b=randint(1,lis[1])
    if a+b>lis[1]:
        continue
    if index%2==0:
        p.add_run("\n").font.size=Pt(32)
    string=f"{a}+{b}"+"="+"\t\t\t\t"
    p.add_run(string).font.size=Pt(32)
    index+=1
doc.save(abs_path+"index.docx")
print(printa.open_file)
printa.print_word(printa.open_file)